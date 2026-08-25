"""wongo.styles — data-driven house-look layer (HANDOFF step 2).

TASTE, not correctness (ground rule 3): everything here shapes how the
document LOOKS and is selected by a style profile YAML under styles/
(default.yml = neutral baseline, kist-wcr.yml = KIST Water Cycle Research
house look). Unconditional OOXML correctness lives in wongo.docxpatch and
runs regardless of the chosen profile.

Resolution order: _journal.yml `style:` key -> $WONGO_STYLE -> `default`.
(Historical note: during the pre-v0.1.0 transition the final fallback was
temporarily `kist-wcr` because the live manuscript could not gain a style
key before ms-r0-sent; that wart was resolved in v0.1.0.)

Style profiles migrated verbatim from the HOUSE_* constants of
legacy/render.py; the WR title-block machinery (rebuild_title_block,
read_front_matter, inject_keywords, bold_caption_leads, _restyle_caption)
moved here parameterized by the profile.
"""
from __future__ import annotations

import os
import re
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from wongo.docxpatch import (
    _style_by_name,
    _tbl_rescale_grid,
    _tbl_set_borders,
    _tbl_set_width_pct,
    dedupe_ppr,
    normalize_ppr_order,
)


DEFAULT_STYLE = "default"


def styles_dir() -> Path:
    """Locate the style-profile directory. $WONGO_STYLES_DIR overrides (tests);
    otherwise the packaged profiles shipped inside this package (wheel-safe)."""
    override = os.environ.get("WONGO_STYLES_DIR")
    if override:
        return Path(override)
    here = Path(__file__).resolve().parent
    if (here / "default.yml").exists():
        return here
    raise SystemExit("packaged style profiles missing; set WONGO_STYLES_DIR")


def load_style(name: str | None) -> dict:
    """Load styles/<name>.yml; None/empty resolves to the neutral default."""
    fname = name or DEFAULT_STYLE
    path = styles_dir() / f"{fname}.yml"
    if not path.exists():
        known = sorted(p.stem for p in styles_dir().glob("*.yml"))
        raise SystemExit(f"style profile '{fname}' not found: {path}. Known: {', '.join(known)}")
    style = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    style["_name"] = fname
    return style


# ---------------------------------------------------------------------------
# Front-matter reading (moved verbatim from legacy/render.py)


def read_front_matter(project: Path, qmd: str = "index.qmd") -> dict:
    """Parse the leading YAML block of a .qmd (pandoc drops author
    affiliations from the docx entirely, so the title-page rebuild needs the
    source metadata)."""
    lines = (project / qmd).read_text(encoding="utf-8").splitlines()
    if not lines or lines[0].strip() != "---":
        return {}
    try:
        end = next(i for i, ln in enumerate(lines[1:], start=1) if ln.strip() == "---")
    except StopIteration:
        return {}
    return yaml.safe_load("\n".join(lines[1:end])) or {}


# ---------------------------------------------------------------------------
# Page shape and paragraph look


def apply_page_geometry(doc: Document, style: dict) -> None:
    page = style.get("page") or {}
    size = page.get("size_mm")
    margins = page.get("margins_mm")
    if not size or not margins:
        return
    w, h = size
    for s in doc.sections:
        s.page_width, s.page_height = Mm(w), Mm(h)
        s.left_margin, s.right_margin = Mm(margins["left"]), Mm(margins["right"])
        s.top_margin, s.bottom_margin = Mm(margins["top"]), Mm(margins["bottom"])


def apply_spacing_and_alignment(doc: Document, style: dict) -> None:
    spacing = style.get("spacing") or {}
    for name in spacing.get("double") or ():
        st = _style_by_name(doc, name)
        if st is not None:
            st.paragraph_format.line_spacing = 2.0
    for name in spacing.get("single") or ():
        st = _style_by_name(doc, name)
        if st is not None:
            st.paragraph_format.line_spacing = 1.0
    for name in style.get("justify") or ():
        st = _style_by_name(doc, name)
        if st is not None:
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def apply_heading_look(doc: Document, style: dict) -> None:
    space_before = style.get("headings_space_before_pt") or {}
    for name, pt in (style.get("headings_pt") or {}).items():
        st = _style_by_name(doc, name)
        if st is None:
            continue
        st.font.size = Pt(pt)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        if space_before.get(name):
            st.paragraph_format.space_before = Pt(space_before[name])


def restyle_title(doc: Document, style: dict) -> None:
    title_pt = style.get("title_pt")
    if not title_pt:
        return
    for p in doc.paragraphs:
        if p.style.name == "Title":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(title_pt)
            return


def restyle_abstract_heading(doc: Document, style: dict) -> None:
    opts = style.get("abstract_title")
    if not opts:
        return
    for p in doc.paragraphs:
        if p.style.name == "Abstract Title":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(opts.get("pt", 12))
            return


# ---------------------------------------------------------------------------
# WR-style title-page rebuild (verbatim logic, gated on title_block.style=="wr")


def _initials(name: str) -> str:
    parts = name.split()
    if len(parts) < 2:
        return name
    return " ".join(f"{p[0]}." for p in parts[:-1]) + " " + parts[-1]


def _compose_affiliation(aff: dict) -> str:
    return ", ".join(
        str(aff[k]) for k in ("name", "address", "country") if aff.get(k)
    )


def rebuild_title_block(doc: Document, meta: dict, opts: dict) -> None:
    """Replace pandoc's one-name-per-line Author paragraphs with the WR-style
    block: one author line with superscript affiliation letters (* marks the
    corresponding author), lettered affiliation lines, then the corresponding
    author's e-mail. All metadata comes from the .qmd YAML."""
    authors = meta.get("author") or []
    if not isinstance(authors, list) or not authors or not isinstance(authors[0], dict):
        return
    paras = doc.paragraphs
    idx = [i for i, p in enumerate(paras) if p.style.name == "Author"]
    if not idx:
        return
    aff_order: list[str] = []
    for a in authors:
        for aff in a.get("affiliations") or []:
            s = _compose_affiliation(aff)
            if s and s not in aff_order:
                aff_order.append(s)
    letters = {s: chr(ord("a") + i) for i, s in enumerate(aff_order)}

    first = paras[idx[0]]
    for r in list(first.runs):
        r._r.getparent().remove(r._r)
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for k, a in enumerate(authors):
        first.add_run(a.get("name", ""))
        sup = ",".join(letters[_compose_affiliation(x)] for x in a.get("affiliations") or [])
        if a.get("corresponding"):
            sup = f"{sup},*" if sup else "*"
        if sup:
            first.add_run(" ")
            sr = first.add_run(sup)
            sr.font.superscript = True
        if k < len(authors) - 1:
            first.add_run(", ")

    anchor = paras[idx[-1] + 1]
    for i in idx[1:]:
        paras[i]._p.getparent().remove(paras[i]._p)

    author_style = first.style
    if opts.get("blank_line_before_authors", True):
        first.insert_paragraph_before("", style=author_style)  # blank line after the title
    for s in aff_order:
        p = anchor.insert_paragraph_before("", style=author_style)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = p.add_run(letters[s])
        sr.font.superscript = True
        p.add_run(" " + s)
    corr = [a for a in authors if a.get("corresponding")] or [
        a for a in authors if a.get("email")
    ]
    if corr:
        if opts.get("blank_line_before_corresponding", True):
            anchor.insert_paragraph_before("", style=author_style)
        label = "* Corresponding author" + ("s" if len(corr) > 1 else "") + ":"
        p = anchor.insert_paragraph_before(label, style=author_style)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        emails = ", ".join(
            f"{a['email']} ({_initials(a.get('name', ''))})" for a in corr if a.get("email")
        )
        p = anchor.insert_paragraph_before("E-mail: " + emails, style=author_style)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if opts.get("blank_line_after_block", True):
        anchor.insert_paragraph_before("", style=author_style)  # blank line before the Abstract


def inject_keywords(doc: Document, meta: dict, enabled: bool = True) -> None:
    if not enabled:
        return
    kws = meta.get("keywords")
    if not kws:
        return
    text = "Keywords: " + (", ".join(map(str, kws)) if isinstance(kws, list) else str(kws))
    abstract_idx = [i for i, p in enumerate(doc.paragraphs) if p.style.name == "Abstract"]
    if not abstract_idx:
        return
    anchor = doc.paragraphs[abstract_idx[-1] + 1]
    style = _style_by_name(doc, "Body Text") or doc.paragraphs[abstract_idx[-1]].style
    p = anchor.insert_paragraph_before(text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ---------------------------------------------------------------------------
# Caption leads (verbatim logic; gated on captions.lead == "bold")


# Quarto (1.10.x) wraps every crossref float (figure or kable table) in a 1x1
# outer table; the caption lands INSIDE that cell as a paragraph styled plain
# "Normal" with an NBSP between the word and the number ("Figure\xa01: ...").
# Body-level "Image Caption"/"Table Caption" styles only appear on unlabeled
# floats. Both shapes are handled below. See docx-quirks.
_CAPTION_LEAD = re.compile(r"^((?:Figure|Table)[\s ]+S?[\s ]?\d+)([.:])[\s ]*(.*)$", re.S)
_CAPTION_LEAD_LOOSE = re.compile(r"^((?:Figure|Table)[\s ]+S?[\s ]?\d+)[.:]?[\s ]*(.*)$", re.S)


def _restyle_caption(p, lead: str, rest: str, opts: dict) -> None:
    # NBSPs -> spaces, then close the gap crossref leaves inside "S 1"
    lead = re.sub(r"S\s+(\d)", r"S\1", lead.replace(" ", " ")) + opts.get("delimiter", ".")
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    br = p.add_run(lead)
    br.font.bold = True
    if rest:
        p.add_run(" " + rest)
    if opts.get("spacing", "single") == "single":
        p.paragraph_format.line_spacing = 1.0
    align = opts.get("align", "justify")
    if align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def bold_caption_leads(doc: Document, opts: dict) -> None:
    """WR-family caption look: bold 'Figure 1.' / 'Table S2.' lead, period
    delimiter regardless of what crossref emitted, single-spaced."""
    # body-level caption-styled paragraphs (unlabeled floats)
    for p in doc.paragraphs:
        if p.style.name in ("Caption", "Image Caption", "Table Caption"):
            m = _CAPTION_LEAD_LOOSE.match(p.text)
            if m:
                _restyle_caption(p, m.group(1), m.group(2), opts)
    # captions inside crossref wrapper-table cells: require the [.:] delimiter
    # so prose like "Table 1 shows..." inside a real data cell never matches
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    m = _CAPTION_LEAD.match(p.text)
                    if m:
                        _restyle_caption(p, m.group(1), m.group(3), opts)


# ---------------------------------------------------------------------------
# Table look (booktabs family; grid rescale itself is unconditional engine
# behavior in wongo.docxpatch — this only decides borders/bold/wrapper rules)


def _bold_cells(cells) -> None:
    for cell in cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True


def _spacer_after(tbl) -> None:
    """One empty single-spaced paragraph after a table so body text doesn't
    sit flush against its bottom rule (a bare empty paragraph would inherit
    Normal's double spacing and gape)."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)
    p.append(pPr)
    tbl._tbl.addnext(p)


def fix_tables(doc: Document, style: dict) -> None:
    """WR/booktabs table look when the profile asks for it. Quarto wraps every
    crossref float in a 1x1 outer table; real data tables live nested inside
    those cells (or stand alone in the SI). Wrappers: borderless, rescaled to
    the full text column. Data tables: full width, closed 1pt top and bottom
    rules (the style's header midrule stays), no vertical/inner rules, bold
    header row and bold first (label) column."""
    sec = doc.sections[0]
    text_w = (int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin)) // 635
    tbl = style.get("tables") or {}

    def walk(tables, width_dxa, top=False):
        for t in tables:
            is_wrapper = len(t.rows) == 1 and len(t.columns) == 1
            if tbl.get("width", "full") == "full":
                _tbl_set_width_pct(t)
                _tbl_rescale_grid(t, width_dxa)
            if is_wrapper:
                if tbl.get("rules") == "booktabs":
                    _tbl_set_borders(t, top=None, bottom=None)
                holds_table = False
                for row in t.rows:
                    for cell in row.cells:
                        if cell.tables:
                            holds_table = True
                        walk(cell.tables, width_dxa - 216)  # minus cell margins
                if top and holds_table and tbl.get("spacer_after"):
                    _spacer_after(t)  # table floats only; figure floats keep their flow
            elif tbl.get("rules") == "booktabs":
                _tbl_set_borders(t, top=8, bottom=8)
                if tbl.get("bold_header_row"):
                    _bold_cells(t.rows[0].cells)
                if tbl.get("bold_first_column"):
                    _bold_cells(row.cells[0] for row in t.rows)
                if top and tbl.get("spacer_after"):
                    _spacer_after(t)

    walk(doc.tables, text_w, top=True)


# ---------------------------------------------------------------------------
# Orchestrator — order of operations is byte-pinned; do not reorder.


def apply_style(doc: Document, style: dict, meta: dict | None = None) -> None:
    """Compose correctness passes (docxpatch) with the selected house look."""
    dedupe_ppr(doc)
    apply_page_geometry(doc, style)
    if (style.get("line_numbers") or "none") == "continuous":
        from wongo.docxpatch import add_line_numbers

        add_line_numbers(doc)
    apply_spacing_and_alignment(doc, style)
    apply_heading_look(doc, style)
    restyle_title(doc, style)
    tb = style.get("title_block") or {}
    if meta:
        if tb.get("style") == "wr":
            rebuild_title_block(doc, meta, tb)
        inject_keywords(doc, meta, enabled=bool(style.get("keywords_line")))
    restyle_abstract_heading(doc, style)
    caps = style.get("captions") or {}
    if caps.get("lead") == "bold":
        bold_caption_leads(doc, caps)
    fix_tables(doc, style)
    normalize_ppr_order(doc)
