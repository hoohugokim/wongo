"""wongo.engine — render pipeline (HANDOFF step 4).

Migrated from legacy/render.py: quarto invocation, target logic, TOC-art
gate, SI pipeline, and post-processing orchestration. Validation checks live
in wongo.engine.checks; tracked-changes extraction in wongo.engine.roundtrip.

TWO LONG-STANDING SI COVER GAPS ARE FIXED HERE (tests in
tests/test_engine.py pin them):
1. the cover sheet printed "Journal: ... — ms_type" where journal profiles
   specify AUTHORS + title (ES&T verified 2026-07-03: cover sheet carries
   authors, title, page/figure/table counts);
2. the page-count line shipped a literal placeholder string — replaced with
   a NUMPAGES field that Word resolves to the true count on open.
"""
from __future__ import annotations

import os
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm

from wongo import styles as wstyles
from wongo.docxpatch import (
    add_line_numbers,
    add_page_numbers,
    patch_theme_fonts,
    restart_page_numbering,
    set_fonts,
)

FONTS = {"collab": "Cambria", "submission": "Cambria"}


def set_line_spacing(doc: Document, factor: float) -> None:
    doc.styles["Normal"].paragraph_format.line_spacing = factor


def resolve_style(cfg: dict) -> dict:
    """Style-profile resolution: _journal.yml `style:` key -> $WONGO_STYLE ->
    `default`. The CLI's --style flag slots in front of this chain; an
    explicit `style: kist-wcr` in _journal.yml (as the reference manuscript now pins) selects
    the KIST-WCR house look."""
    name = cfg.get("style") or os.environ.get("WONGO_STYLE") or "default"
    return wstyles.load_style(name)


# ---------------------------------------------------------------------------
# SI cover sheet


def _numpages_field() -> OxmlElement:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " NUMPAGES ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    return fld


def prepend_si_cover(doc: Document, profile: dict, counts: dict, meta: dict | None = None) -> None:
    """Cover sheet per journal profile: title, authors, figure/table counts,
    and a page count Word fills in on open (python-docx cannot paginate)."""
    meta = meta or {}
    lines = ["Supporting Information"]
    if meta.get("title"):
        lines.append(str(meta["title"]))
    authors = [
        a.get("name", "") for a in (meta.get("author") or []) if isinstance(a, dict)
    ]
    if authors:
        lines.append(", ".join(n for n in authors if n))
    lines.append(f"Contents: {counts.get('figures', 0)} figures, {counts.get('tables', 0)} tables")

    first = doc.paragraphs[0]
    for text in lines:
        p = first.insert_paragraph_before(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pages = first.insert_paragraph_before("")
    pages.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pages.add_run("Pages: ")
    pages._p.append(_numpages_field())
    brk = first.insert_paragraph_before("")
    run = brk.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


DEFAULT_TOC_ART_LABEL = "For Table of Contents Only"


def insert_toc_art(
    doc: Document,
    image_path: Path,
    width_mm: float | None = None,
    label: str = DEFAULT_TOC_ART_LABEL,
) -> None:
    """Append the TOC/abstract graphic at the END of the document.

    Per the official ACS "Guidelines for Table of Contents/Abstract
    Graphics" (dated 2024-02-28): the graphic is labeled "For Table of
    Contents Only" and placed on the LAST PAGE of the submitted manuscript
    (it may also be uploaded separately as "Graphics for manuscript").
    Produces, in order: a page break, a centered label paragraph, then a
    centered paragraph containing the image.
    """
    brk = doc.add_paragraph()
    run = brk.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)

    label_p = doc.add_paragraph(label)
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    kwargs = {"width": Mm(width_mm)} if width_mm else {}
    run.add_picture(str(image_path), **kwargs)


# ---------------------------------------------------------------------------
# Quarto invocation + gates


def quarto_render(project: Path, qmd: str, out_name: str, profile: dict) -> Path:
    cmd = ["quarto", "render", qmd, "--to", "docx", "--output", out_name]
    ref = profile.get("reference_doc")
    if ref:
        cmd += ["-M", f"reference-doc:{Path(profile['_dir']) / ref}"]
    csl = profile.get("csl")
    if csl:
        cmd += ["-M", f"csl:{Path(profile['_dir']) / csl}"]
    try:
        subprocess.run(cmd, cwd=project, check=True)
    except subprocess.CalledProcessError:
        raise SystemExit(f"quarto render failed for {qmd} (see quarto output above)")
    out_dir = project / "output"
    out_dir.mkdir(exist_ok=True)
    produced = project / out_name          # quarto writes next to the input
    target = out_dir / out_name
    if produced.exists() and produced != target:
        shutil.move(str(produced), target)
    if not target.exists():
        raise SystemExit(f"quarto render did not produce {target}")
    return target


def find_toc_art(project: Path) -> Path | None:
    for ext in ("png", "tiff", "tif", "jpg", "jpeg"):
        p = project / "figures" / f"toc-art.{ext}"
        if p.exists():
            return p
    return None


# ---------------------------------------------------------------------------
# Post-processing


def postprocess_main(path: Path, profile: dict, cfg: dict, target: str, project: Path) -> None:
    doc = Document(str(path))
    if target == "submission":
        if profile.get("line_numbers"):
            add_line_numbers(doc)
        if profile.get("spacing") == "double":
            set_line_spacing(doc, 2.0)
        toc = (profile.get("toc_graphic") or {})
        if toc.get("required"):
            art = find_toc_art(project)
            if art is None:
                # Backstop only: render_project() gates this BEFORE
                # quarto_render() runs, so this branch should be unreachable in
                # practice. Message kept aligned with the pre-render gate.
                raise SystemExit(
                    "Profile requires TOC art but figures/toc-art.{png,tif,tiff,jpg,jpeg} is missing."
                )
            insert_toc_art(doc, art, toc.get("width_mm"), toc.get("label") or DEFAULT_TOC_ART_LABEL)
    style = resolve_style(cfg)
    font = style.get("font") or FONTS[target]
    set_fonts(doc, font)
    wstyles.apply_style(doc, style, wstyles.read_front_matter(project, "index.qmd"))
    add_page_numbers(doc)
    doc.save(str(path))
    patch_theme_fonts(path, font, track_changes=(target == "collab"))


def postprocess_si(path: Path, profile: dict, cfg: dict, target: str, project: Path) -> None:
    doc = Document(str(path))
    si = profile.get("si") or {}
    style = resolve_style(cfg)
    font = style.get("font") or FONTS[target]
    set_fonts(doc, font)
    # no index.qmd meta for styling: SI keeps its plain author line, no keywords;
    # but the cover sheet DOES want title/authors from the front matter
    meta = wstyles.read_front_matter(project, "index.qmd") if si.get("needs_cover_sheet") else None
    wstyles.apply_style(doc, style)
    counts = {
        "figures": len(doc.inline_shapes),
        "tables": len(doc.tables),
    }
    if si.get("needs_cover_sheet"):
        prepend_si_cover(doc, profile, counts, meta)
    add_page_numbers(doc, prefix=si.get("page_prefix", "S"))
    restart_page_numbering(doc, start=1)
    doc.save(str(path))
    patch_theme_fonts(path, font, track_changes=(target == "collab"))


# ---------------------------------------------------------------------------
# Orchestration


def check_hard_failures(project: Path) -> list:
    """Run validation checks; returns the HARD failures ([] if clean)."""
    from wongo.engine.checks import run_checks

    checks = run_checks(project)
    from wongo.engine.checks import print_report

    print_report(checks)
    return [c for c in checks if c.level == "HARD" and not c.ok]


def render_project(project: Path, target: str) -> int:
    """The full render pipeline for one manuscript project. Returns 0 on
    success; raises SystemExit on gate failures."""
    from wongo.profiles import load_journal_config, load_profile, manuscript_type

    project = Path(project).resolve()
    cfg = load_journal_config(project)
    profile = load_profile(cfg["journal"], project)
    manuscript_type(profile, cfg["ms_type"])  # fail fast on bad ms_type

    hard = check_hard_failures(project)
    if hard and target == "submission":
        raise SystemExit("HARD checks failed — submission render refused (fix, or render --target collab).")

    # Gate BEFORE quarto ever runs: refusing here means no half-processed
    # deliverable is left behind when the TOC art is missing (backstop kept
    # in postprocess_main in case this function is ever bypassed).
    if target == "submission" and (profile.get("toc_graphic") or {}).get("required"):
        if find_toc_art(project) is None:
            raise SystemExit(
                "Profile requires TOC art but figures/toc-art.{png,tif,tiff,jpg,jpeg} "
                "is missing — submission render refused."
            )

    main_docx = quarto_render(project, "index.qmd", f"main-{target}.docx", profile)
    postprocess_main(main_docx, profile, cfg, target, project)
    print(f"wrote {main_docx}")

    if (project / "si.qmd").exists():
        si_docx = quarto_render(project, "si.qmd", f"si-{target}.docx", profile)
        postprocess_si(si_docx, profile, cfg, target, project)
        print(f"wrote {si_docx}")
    return 0
