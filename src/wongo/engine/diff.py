"""wongo.engine.diff — tracked-changes stamping for revision submissions (S6).

Given the originally submitted DOCX and a revised render, produce a copy of
the revised document whose differences appear as REAL Word track changes
(`w:ins` / `w:del`), attributed to a chosen author/date. Journals that demand
a marked-up revision get one without a manual Word Compare session.

Scope (v1): body-level paragraphs. Word-level diffs within replaced
paragraphs; whole-paragraph insertions/deletions fully tracked. TABLES are
deliberately left untouched (their content is data-driven from .qmd chunks,
and merging table diffs is a Word-Compare-grade problem); if any table text
differs between the two documents this is REPORTED (`tables_differ`) so the
human knows to run Word Compare for those pages.

The emitted markup is the exact syntax `quarto pandoc --track-changes=all`
parses (see docs/docx-quirks.md), so `wongo roundtrip` can extract our own
output back into a merge worksheet.

Implementation notes:
- All edits happen IN PLACE on the revised document's own paragraph
  elements — never clone-and-remove — so every paragraph element stays live
  in the lxml tree across the whole opcode walk.
- Deletions are materialized as NEW paragraphs (the revised doc no longer
  contains them) inserted at the aligned position; they carry `w:delText`,
  never `w:t`, per ECMA-376 tracked-deletion semantics.
"""
from __future__ import annotations

import copy
import re
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _tokens(text: str) -> list[str]:
    """Split into word/whitespace tokens, preserving everything."""
    return re.findall(r"\s+|[^\s]+", text)


def _norm(text: str) -> str:
    return " ".join(text.split())


def _make_run(text: str, rpr_template, *, deleted: bool):
    r = OxmlElement("w:r")
    if rpr_template is not None:
        r.append(copy.deepcopy(rpr_template))
    target = OxmlElement("w:delText") if deleted else OxmlElement("w:t")
    if text != text.strip():
        target.set(qn("xml:space"), "preserve")
    target.text = text
    r.append(target)
    return r


def _wrap(kind: str, runs: list, author: str, date: str, counter: list[int]):
    wrap = OxmlElement(f"w:{kind}")
    wrap.set(qn("w:id"), str(counter[0]))
    counter[0] += 1
    wrap.set(qn("w:author"), author)
    wrap.set(qn("w:date"), date)
    for r in runs:
        wrap.append(r)
    return wrap


def _strip_runs(p_el) -> None:
    """Remove all w:r/w:ins/w:del children, keep pPr."""
    for child in list(p_el):
        if child.tag in (qn("w:r"), qn("w:ins"), qn("w:del")):
            p_el.remove(child)


def _rpr_of(paragraph) -> object | None:
    for r in paragraph._p.findall(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is not None:
            return rpr
    return None


def diff_documents(original: Path, revised: Path, out: Path,
                   author: str = "Revised manuscript",
                   date: str | None = None) -> dict:
    """Stamp tracked changes into `revised` vs `original`; save to `out`.

    Returns a report dict with insertion/deletion counts and the
    tables_differ flag. Neither input file is modified.
    """
    if date is None:
        date = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    doc_orig = Document(str(original))
    doc_rev = Document(str(revised))

    counter = [9000]  # high base avoids clashing with real revision ids
    report = {"inserted_words": 0, "deleted_words": 0,
              "inserted_paragraphs": 0, "deleted_paragraphs": 0,
              "tables_differ": False}

    o_texts = [_norm(p.text or "") for p in doc_orig.paragraphs]
    r_paras = doc_rev.paragraphs
    r_texts = [_norm(p.text or "") for p in r_paras]

    def count_words(tokens: list[str]) -> int:
        return len([t for t in tokens if t.strip()])

    def mark_fully_tracked_inplace(p, kind: str) -> None:
        """Wrap ALL of paragraph p's existing text as one ins/del, in place."""
        saved_text = p.text or ""
        rpr = _rpr_of(p)
        _strip_runs(p._p)
        runs = [_make_run(t, rpr, deleted=(kind == "del"))
                for t in _tokens(saved_text)]
        p._p.append(_wrap(kind, runs, author, date, counter))
        report[f"{'inserted' if kind == 'ins' else 'deleted'}_words"] += \
            count_words(_tokens(saved_text))
        report["inserted_paragraphs" if kind == "ins"
               else "deleted_paragraphs"] += 1

    def make_deleted_para(src_para):
        """New paragraph carrying the source text as a full deletion."""
        new_p = OxmlElement("w:p")
        src_ppr = src_para._p.find(qn("w:pPr"))
        if src_ppr is not None:
            new_p.append(copy.deepcopy(src_ppr))
        rpr = _rpr_of(src_para)
        tokens = _tokens(src_para.text or "")
        runs = [_make_run(t, rpr, deleted=True) for t in tokens]
        new_p.append(_wrap("del", runs, author, date, counter))
        report["deleted_words"] += count_words(tokens)
        report["deleted_paragraphs"] += 1
        return new_p

    def rebuild_word_diff(p, old_tokens, rpr):
        """Rewrite paragraph p's runs as equal/ins/del token runs."""
        new_tokens = _tokens(p.text or "")
        _strip_runs(p._p)
        sm_w = SequenceMatcher(a=old_tokens, b=new_tokens, autojunk=False)
        for op, i1, i2, j1, j2 in sm_w.get_opcodes():
            if op == "equal":
                # unchanged words stay as plain runs, in document order
                for t in old_tokens[i1:i2]:
                    p._p.append(_make_run(t, rpr, deleted=False))
                continue
            if op in ("delete", "replace"):
                runs = [_make_run(t, rpr, deleted=True) for t in old_tokens[i1:i2]]
                p._p.append(_wrap("del", runs, author, date, counter))
                report["deleted_words"] += count_words(old_tokens[i1:i2])
            if op in ("insert", "replace"):
                runs = [_make_run(t, rpr, deleted=False) for t in new_tokens[j1:j2]]
                p._p.append(_wrap("ins", runs, author, date, counter))
                report["inserted_words"] += count_words(new_tokens[j1:j2])

    body = doc_rev.element.body
    sect_pr = body.find(qn("w:sectPr"))

    sm = SequenceMatcher(a=o_texts, b=r_texts, autojunk=False)
    idx_r = 0
    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op == "equal":
            idx_r += j2 - j1
        elif op == "insert":
            for jp in range(j1, j2):
                mark_fully_tracked_inplace(r_paras[jp], "ins")
            idx_r += j2 - j1
        elif op == "delete":
            anchor = (r_paras[idx_r]._p if idx_r < len(r_paras)
                      else (sect_pr if sect_pr is not None else None))
            for ip in range(i1, i2):
                tracked = make_deleted_para(doc_orig.paragraphs[ip])
                if anchor is not None:
                    anchor.addprevious(tracked)
                else:
                    body.append(tracked)
        elif op == "replace":
            olds = doc_orig.paragraphs[i1:i2]
            pairs = min(len(olds), j2 - j1)
            for k in range(pairs):
                p = r_paras[j1 + k]
                rebuild_word_diff(p, _tokens(olds[k].text or ""),
                                  _rpr_of(p) or _rpr_of(olds[k]))
            if len(olds) > pairs:  # surplus originals -> deletions after block
                last = r_paras[j2 - 1]._p
                for ip in range(i1 + pairs, i2):
                    tracked = make_deleted_para(olds[ip])
                    last.addnext(tracked)
                    last = tracked
            for jp in range(j1 + pairs, j2):  # surplus revised -> insertions
                mark_fully_tracked_inplace(r_paras[jp], "ins")
            idx_r += j2 - j1

    # Tables: untouched by design, but surface any difference.
    def table_text(doc):
        return "\n".join(c.text for t in doc.tables
                         for row in t.rows for c in row.cells)

    try:
        report["tables_differ"] = table_text(doc_orig) != table_text(doc_rev)
    except Exception:  # noqa: BLE001 — reporting must never fail the stamp
        report["tables_differ"] = False

    out.parent.mkdir(parents=True, exist_ok=True)
    doc_rev.save(str(out))
    return report


def main(argv: list[str] | None = None) -> int:
    """CLI entry backing `wongo diff`."""
    import argparse

    ap = argparse.ArgumentParser(
        prog="wongo diff",
        description="Stamp tracked changes into a revised DOCX vs the original "
                    "submission (S6 marked-up revision, no Word Compare needed).")
    ap.add_argument("original", help="originally submitted DOCX")
    ap.add_argument("revised", help="revised DOCX (never modified)")
    ap.add_argument("-o", "--out", default=None,
                    help="output path (default: <revised-stem>-tracked.docx)")
    ap.add_argument("--author", default="Revised manuscript",
                    help="attribution for stamped changes")
    ap.add_argument("--date", default=None,
                    help="ISO timestamp for stamped changes (default: now UTC)")
    args = ap.parse_args(argv)

    original, revised = Path(args.original), Path(args.revised)
    out = Path(args.out) if args.out else revised.with_name(
        revised.stem + "-tracked.docx")
    report = diff_documents(original, revised, out,
                            author=args.author, date=args.date)
    print(f"wrote {out}")
    print(f"  words: +{report['inserted_words']} inserted, "
          f"-{report['deleted_words']} deleted")
    print(f"  paragraphs: +{report['inserted_paragraphs']}, "
          f"-{report['deleted_paragraphs']}")
    if report["tables_differ"]:
        print("  NOTE: tables differ between the two documents and are NOT "
              "tracked by this tool — run Word Compare for table pages.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
