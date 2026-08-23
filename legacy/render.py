#!/usr/bin/env python3
# /// script
# requires-python = ">=3.11"
# dependencies = ["PyYAML>=6", "python-docx>=1.1"]
# ///
"""Render a Quarto manuscript to collaboration- or submission-grade DOCX.

Usage: render.py --target {collab|submission} [--project DIR]

Pipeline: validate gate (submission refuses on HARD failures) -> quarto render
with the journal profile's reference-doc/CSL -> python-docx post-processing
driven by profile.yml (line numbers, spacing, fonts, page numbers, TOC art) ->
separate SI render with S-page numbering and cover sheet.
Fonts: both targets use Cambria (Cambria Math is the MS Word default for math).
Partial-state note: if the SI render fails after the main render already
succeeded, output/main-<target>.docx from this run remains on disk — rerun
after fixing the SI source; the main render is not repeated for you.
"""
from __future__ import annotations

import argparse
import shutil
import subprocess
from pathlib import Path

import re

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

import mslib
import validate

FONTS = {"collab": "Cambria", "submission": "Cambria"}
STYLES_TO_FONT = (
    "Normal", "Title", "Subtitle", "Author", "Date", "Abstract", "Abstract Title",
    "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6",
    "Caption", "Image Caption", "Table Caption", "First Paragraph", "Body Text",
    "Bibliography", "Compact", "Captioned Figure", "Block Text",
)


# CT_SectPr child sequence per ECMA-376: ... lnNumType, pgNumType, then these.
# w:lnNumType's successors are therefore ("w:pgNumType",) + _SECT_PR_TAIL;
# w:pgNumType's are _SECT_PR_TAIL. See quarto-docx-quirks.md for why plain
# .append() corrupts real journal templates.
_SECT_PR_TAIL = (
    "w:cols", "w:formProt", "w:vAlign", "w:noEndnote", "w:titlePg",
    "w:textDirection", "w:bidi", "w:rtlGutter", "w:docGrid",
    "w:printerSettings", "w:sectPrChange",
)


def add_line_numbers(doc: Document) -> None:
    for section in doc.sections:
        sect_pr = section._sectPr
        ln = sect_pr.find(qn("w:lnNumType"))
        if ln is None:
            ln = OxmlElement("w:lnNumType")
            sect_pr.insert_element_before(ln, *(("w:pgNumType",) + _SECT_PR_TAIL))
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:restart"), "continuous")


def set_line_spacing(doc: Document, factor: float) -> None:
    doc.styles["Normal"].paragraph_format.line_spacing = factor


def _style_by_name(doc: Document, style_name: str):
    """Find a style by its exact UI name.

    Deliberately avoids ``doc.styles[style_name]``: python-docx's dict-style
    lookup applies a "BabelFish" UI<->internal name translation (e.g.
    "Caption" -> "caption") before matching by name, and only falls back to a
    *deprecated* style-id match if that fails. Quarto/pandoc's default
    reference.docx stores the built-in Caption style's internal XML name as
    "Caption" (capitalized) rather than Word's own "caption" (lowercase), so
    the translated lookup misses and every call silently takes the
    deprecated path, emitting a UserWarning today and liable to raise
    KeyError once python-docx removes that fallback. Iterating styles and
    comparing `.name` directly sidesteps the translation entirely.
    """
    for style in doc.styles:
        if style.name == style_name:
            return style
    return None


def set_fonts(doc: Document, name: str) -> None:
    for style_name in STYLES_TO_FONT:
        style = _style_by_name(doc, style_name)
        if style is None:
            continue
        style.font.name = name
        rfonts = style.element.get_or_add_rPr().get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), name)
        rfonts.set(qn("w:cs"), name)
        # OOXML precedence trap: a w:asciiTheme/w:hAnsiTheme attribute BEATS the
        # literal w:ascii/w:hAnsi on the same rFonts, so pandoc's theme-linked
        # heading styles keep rendering in the document theme's font (Aptos in
        # current Word) no matter what font.name is set to. Strip the theme
        # attributes so the literal font wins. See quarto-docx-quirks.md.
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            if rfonts.get(qn(attr)) is not None:
                del rfonts.attrib[qn(attr)]


def patch_theme_fonts(path: Path, name: str, track_changes: bool = False) -> None:
    """Zip-level fixes python-docx can't do: (1) rewrite the document theme's
    typefaces so anything resolving through the theme lands on the same font
    instead of Word's Aptos default; (2) stamp compatibilityMode 15 into
    settings.xml — pandoc emits NO compatSetting at all, so Word opens every
    render in Compatibility Mode and applies legacy layout rules (pct table
    widths and in-cell justification misbehave); (3) when track_changes is
    True (collab target only — NEVER submission), plant <w:trackChanges/> so
    the document OPENS with Track Changes already on: coauthors start editing
    and every edit is captured without anyone remembering the toggle. It is a
    default, not a lock — a reviewer can still switch it off; we deliberately
    skip documentProtection (password-hash lock) as hostile to coauthors.
    See quarto-docx-quirks.md."""
    import zipfile

    with zipfile.ZipFile(str(path)) as z:
        names = z.namelist()
        if "word/theme/theme1.xml" not in names:
            return
        items = {n: z.read(n) for n in names}
    theme = items["word/theme/theme1.xml"].decode("utf-8")
    theme = re.sub(
        r'(<a:(?:latin|ea|cs)\s+typeface=")[^"]*(")',
        lambda m: m.group(1) + name + m.group(2),
        theme,
    )
    items["word/theme/theme1.xml"] = theme.encode("utf-8")
    settings = items.get("word/settings.xml", b"").decode("utf-8")
    if settings and "compatibilityMode" not in settings:
        compat = ('<w:compat><w:compatSetting w:name="compatibilityMode" '
                  'w:uri="http://schemas.microsoft.com/office/word" w:val="15"/></w:compat>')
        if "<w:rsids" in settings:
            settings = settings.replace("<w:rsids", compat + "<w:rsids", 1)
        else:
            settings = settings.replace("</w:settings>", compat + "</w:settings>", 1)
        items["word/settings.xml"] = settings.encode("utf-8")
    elif settings:
        # present but possibly < 15 (python-docx's own template ships 14, which
        # ALSO opens in Compatibility Mode) — upgrade in place
        settings = re.sub(
            r'<w:compatSetting[^>]*w:name="compatibilityMode"[^>]*/>',
            lambda m: re.sub(r'w:val="\d+"', 'w:val="15"', m.group(0)),
            settings,
        )
        items["word/settings.xml"] = settings.encode("utf-8")
    settings = items.get("word/settings.xml", b"").decode("utf-8")
    if track_changes and settings and "<w:trackChanges" not in settings:
        # CT_Settings order: trackChanges precedes doNotTrackMoves and
        # defaultTabStop — anchor on whichever exists (pandoc has both).
        el = "<w:trackChanges/>"
        for anchor in ("<w:doNotTrackMoves", "<w:defaultTabStop"):
            if anchor in settings:
                settings = settings.replace(anchor, el + anchor, 1)
                break
        else:
            settings = re.sub(r"(<w:settings[^>]*>)", r"\1" + el, settings, count=1)
        items["word/settings.xml"] = settings.encode("utf-8")
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(str(tmp), "w", zipfile.ZIP_DEFLATED) as z:
        for n, data in items.items():
            z.writestr(n, data)
    tmp.replace(path)


# ---------------------------------------------------------------------------
# HOUSE STYLE (lab preference, venue-neutral). Makes every DOCX product carry
# the page shape, line numbering, double spacing, title-page block and caption
# look of the group's prior Water Research / WR X submissions (see
# manuscript/prev-ref/ in the the reference manuscript project), so coauthors meet a familiar
# page and don't reject the Quarto workflow on looks. Fonts stay whatever
# FONTS says (currently Cambria). Nothing here contradicts verified ES&T
# guidance: line numbers are optional there, spacing is unspecified, and
# "Abstract and Keywords" is an ACS heading. Applied to BOTH targets.
HOUSE_PAGE = {"w": 210, "h": 297, "left": 25, "right": 25, "top": 30, "bottom": 25}
HOUSE_DOUBLE = ("Normal", "Body Text", "First Paragraph", "Abstract")
HOUSE_SINGLE = (  # explicit shield against inheriting Normal's 2.0
    "Compact", "Caption", "Image Caption", "Table Caption", "Footnote Text",
    "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title", "Author", "Date",
)
HOUSE_JUSTIFY = ("Body Text", "First Paragraph", "Abstract", "Bibliography",
                 "Caption", "Image Caption", "Table Caption",
                 "Heading 1", "Heading 2", "Heading 3", "Heading 4")
HOUSE_HEADING_PT = {"Heading 1": 16, "Heading 2": 13, "Heading 3": 12, "Heading 4": 12}
HOUSE_TITLE_PT = 14


def read_front_matter(project: Path, qmd: str = "index.qmd") -> dict:
    """Parse the leading YAML block of a .qmd (pandoc drops author
    affiliations from the docx entirely, so the title-page rebuild needs the
    source metadata)."""
    lines = (project / qmd).read_text(encoding="utf-8").splitlines()
    if not lines or lines[0].strip() != "---":
        return {}
    try:
        end = next(i for i, ln in enumerate(lines[1:], start=1) if ln.strip() == "---")
    except StopIteration:
        return {}
    return yaml.safe_load("\n".join(lines[1:end])) or {}


def apply_page_geometry(doc: Document) -> None:
    for s in doc.sections:
        s.page_width, s.page_height = Mm(HOUSE_PAGE["w"]), Mm(HOUSE_PAGE["h"])
        s.left_margin, s.right_margin = Mm(HOUSE_PAGE["left"]), Mm(HOUSE_PAGE["right"])
        s.top_margin, s.bottom_margin = Mm(HOUSE_PAGE["top"]), Mm(HOUSE_PAGE["bottom"])


def apply_house_spacing_and_alignment(doc: Document) -> None:
    for name in HOUSE_DOUBLE:
        st = _style_by_name(doc, name)
        if st is not None:
            st.paragraph_format.line_spacing = 2.0
    for name in HOUSE_SINGLE:
        st = _style_by_name(doc, name)
        if st is not None:
            st.paragraph_format.line_spacing = 1.0
    for name in HOUSE_JUSTIFY:
        st = _style_by_name(doc, name)
        if st is not None:
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def apply_heading_look(doc: Document) -> None:
    for name, pt in HOUSE_HEADING_PT.items():
        st = _style_by_name(doc, name)
        if st is None:
            continue
        st.font.size = Pt(pt)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)


def restyle_title(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.style.name == "Title":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(HOUSE_TITLE_PT)
            return


def _initials(name: str) -> str:
    parts = name.split()
    if len(parts) < 2:
        return name
    return " ".join(f"{p[0]}." for p in parts[:-1]) + " " + parts[-1]


def _compose_affiliation(aff: dict) -> str:
    return ", ".join(
        str(aff[k]) for k in ("name", "address", "country") if aff.get(k)
    )


def rebuild_title_block(doc: Document, meta: dict) -> None:
    """Replace pandoc's one-name-per-line Author paragraphs with the WR-style
    block: one author line with superscript affiliation letters (* marks the
    corresponding author), lettered affiliation lines, then the corresponding
    author's e-mail. All metadata comes from the .qmd YAML."""
    authors = meta.get("author") or []
    if not isinstance(authors, list) or not authors or not isinstance(authors[0], dict):
        return
    paras = doc.paragraphs
    idx = [i for i, p in enumerate(paras) if p.style.name == "Author"]
    if not idx:
        return
    aff_order: list[str] = []
    for a in authors:
        for aff in a.get("affiliations") or []:
            s = _compose_affiliation(aff)
            if s and s not in aff_order:
                aff_order.append(s)
    letters = {s: chr(ord("a") + i) for i, s in enumerate(aff_order)}

    first = paras[idx[0]]
    for r in list(first.runs):
        r._r.getparent().remove(r._r)
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for k, a in enumerate(authors):
        first.add_run(a.get("name", ""))
        sup = ",".join(letters[_compose_affiliation(x)] for x in a.get("affiliations") or [])
        if a.get("corresponding"):
            sup = f"{sup},*" if sup else "*"
        if sup:
            first.add_run(" ")
            sr = first.add_run(sup)
            sr.font.superscript = True
        if k < len(authors) - 1:
            first.add_run(", ")

    anchor = paras[idx[-1] + 1]
    for i in idx[1:]:
        paras[i]._p.getparent().remove(paras[i]._p)

    author_style = first.style
    first.insert_paragraph_before("", style=author_style)  # blank line after the title
    for s in aff_order:
        p = anchor.insert_paragraph_before("", style=author_style)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = p.add_run(letters[s])
        sr.font.superscript = True
        p.add_run(" " + s)
    corr = [a for a in authors if a.get("corresponding")] or [
        a for a in authors if a.get("email")
    ]
    if corr:
        label = "* Corresponding author" + ("s" if len(corr) > 1 else "") + ":"
        p = anchor.insert_paragraph_before(label, style=author_style)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        emails = ", ".join(
            f"{a['email']} ({_initials(a.get('name', ''))})" for a in corr if a.get("email")
        )
        p = anchor.insert_paragraph_before("E-mail: " + emails, style=author_style)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    anchor.insert_paragraph_before("", style=author_style)  # blank line before the Abstract


def restyle_abstract_heading(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.style.name == "Abstract Title":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(12)
            return


def inject_keywords(doc: Document, meta: dict) -> None:
    kws = meta.get("keywords")
    if not kws:
        return
    text = "Keywords: " + (", ".join(map(str, kws)) if isinstance(kws, list) else str(kws))
    abstract_idx = [i for i, p in enumerate(doc.paragraphs) if p.style.name == "Abstract"]
    if not abstract_idx:
        return
    anchor = doc.paragraphs[abstract_idx[-1] + 1]
    style = _style_by_name(doc, "Body Text") or doc.paragraphs[abstract_idx[-1]].style
    p = anchor.insert_paragraph_before(text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


# Quarto (1.10.x) wraps every crossref float (figure or kable table) in a 1x1
# outer table; the caption lands INSIDE that cell as a paragraph styled plain
# "Normal" with an NBSP between the word and the number ("Figure\xa01: ...").
# Body-level "Image Caption"/"Table Caption" styles only appear on unlabeled
# floats. Both shapes are handled below. See quarto-docx-quirks.md.
_CAPTION_LEAD = re.compile(r"^((?:Figure|Table)[\s ]+S?[\s ]?\d+)([.:])[\s ]*(.*)$", re.S)
_CAPTION_LEAD_LOOSE = re.compile(r"^((?:Figure|Table)[\s ]+S?[\s ]?\d+)[.:]?[\s ]*(.*)$", re.S)


def _restyle_caption(p, lead: str, rest: str) -> None:
    # NBSPs -> spaces, then close the gap crossref leaves inside "S 1"
    lead = re.sub(r"S\s+(\d)", r"S\1", lead.replace(" ", " ")) + "."
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    br = p.add_run(lead)
    br.font.bold = True
    if rest:
        p.add_run(" " + rest)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def bold_caption_leads(doc: Document) -> None:
    """WR-family caption look: bold 'Figure 1.' / 'Table S2.' lead, period
    delimiter regardless of what crossref emitted, single-spaced."""
    # body-level caption-styled paragraphs (unlabeled floats)
    for p in doc.paragraphs:
        if p.style.name in ("Caption", "Image Caption", "Table Caption"):
            m = _CAPTION_LEAD_LOOSE.match(p.text)
            if m:
                _restyle_caption(p, m.group(1), m.group(2))
    # captions inside crossref wrapper-table cells: require the [.:] delimiter
    # so prose like "Table 1 shows..." inside a real data cell never matches
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    m = _CAPTION_LEAD.match(p.text)
                    if m:
                        _restyle_caption(p, m.group(1), m.group(3))


def _tbl_set_width_pct(tbl, pct: int = 5000) -> None:
    tblPr = tbl._tbl.tblPr
    for el in tblPr.findall(qn("w:tblW")):
        tblPr.remove(el)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(pct))
    w.set(qn("w:type"), "pct")
    ts = tblPr.find(qn("w:tblStyle"))
    tblPr.insert(list(tblPr).index(ts) + 1 if ts is not None else 0, w)


def _tbl_set_borders(tbl, top: int | None, bottom: int | None) -> None:
    """Explicit instance-level tblBorders: top/bottom single rules (sz in
    eighths of a point) or none; left/right/inside always none. The Table
    style's conditional firstRow bottom rule (the header midrule) survives,
    because cell-level conditional formatting outranks table-level borders."""
    tblPr = tbl._tbl.tblPr
    for el in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(el)
    borders = OxmlElement("w:tblBorders")
    for edge, sz in (("top", top), ("left", None), ("bottom", bottom),
                     ("right", None), ("insideH", None), ("insideV", None)):
        e = OxmlElement(f"w:{edge}")
        if sz:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(sz))
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), "auto")
        else:
            e.set(qn("w:val"), "none")
        borders.append(e)
    anchor = tblPr.find(qn("w:tblW"))
    tblPr.insert(list(tblPr).index(anchor) + 1 if anchor is not None else 0, borders)


def _tbl_rescale_grid(t, width_dxa: int) -> None:
    """Rescale tblGrid columns (and per-cell tcW) proportionally so they sum
    to width_dxa. Pandoc hard-codes grids for its Letter-width reference doc
    (7920 dxa), so on our A4/25 mm page every table stops ~2 cm short of the
    right margin — and in-cell caption justification then LOOKS left-aligned
    because the cell itself is narrow."""
    grid = t._tbl.find(qn("w:tblGrid"))
    if grid is None:
        return
    cols = grid.findall(qn("w:gridCol"))
    widths = [int(c.get(qn("w:w")) or 0) for c in cols]
    total = sum(widths)
    if not total:
        return
    scaled = [max(1, round(w * width_dxa / total)) for w in widths]
    scaled[-1] += width_dxa - sum(scaled)  # rounding remainder to last col
    for c, w in zip(cols, scaled):
        c.set(qn("w:w"), str(w))
    for row in t.rows:
        for cell, w in zip(row.cells, scaled):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")


def _bold_cells(cells) -> None:
    for cell in cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True


def fix_tables(doc: Document) -> None:
    """WR/booktabs table look. Quarto wraps every crossref float in a 1x1
    outer table; real data tables live nested inside those cells (or stand
    alone in the SI). Wrappers: borderless, rescaled to the full text column.
    Data tables: full width, closed 1pt top and bottom rules (the style's
    header midrule stays), no vertical/inner rules, bold header row and bold
    first (label) column."""
    sec = doc.sections[0]
    text_w = (int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin)) // 635

    def walk(tables, width_dxa):
        for t in tables:
            is_wrapper = len(t.rows) == 1 and len(t.columns) == 1
            _tbl_set_width_pct(t)
            _tbl_rescale_grid(t, width_dxa)
            if is_wrapper:
                _tbl_set_borders(t, top=None, bottom=None)
                for row in t.rows:
                    for cell in row.cells:
                        walk(cell.tables, width_dxa - 216)  # minus cell margins
            else:
                _tbl_set_borders(t, top=8, bottom=8)
                _bold_cells(t.rows[0].cells)
                _bold_cells(row.cells[0] for row in t.rows)
    walk(doc.tables, text_w)


def dedupe_ppr(doc: Document) -> None:
    """Quarto 1.10's docx writer emits a SECOND <w:pPr> on crossref float
    caption paragraphs (and possibly others). That is invalid OOXML: python-docx
    reads/edits the FIRST pPr while Word honors the LAST, so alignment/spacing
    set through python-docx silently never renders. Merge duplicates into one
    pPr (later wins per property, matching what Word displayed), with pStyle
    reordered first as CT_PPr requires. Must run BEFORE any restyling."""
    for para in doc.element.body.iter(qn("w:p")):
        pprs = [c for c in para if c.tag == qn("w:pPr")]
        if len(pprs) < 2:
            continue
        base = pprs[0]
        for extra in pprs[1:]:
            for child in list(extra):
                prev = base.find(child.tag)
                if prev is not None:
                    base.remove(prev)
                base.append(child)
            para.remove(extra)
        st = base.find(qn("w:pStyle"))
        if st is not None:
            base.remove(st)
            base.insert(0, st)


# CT_PPr child sequence per ECMA-376 §17.3.1.26 — Word DISCARDS properties that
# appear out of this order (observed: a jc placed before spacing renders as if
# absent, so "justified" captions stayed left-aligned through two prior fixes).
_PPR_ORDER = (
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
    "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs",
    "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
    "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
    "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents",
    "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr",
    "pPrChange",
)


def normalize_ppr_order(doc: Document) -> None:
    """Re-sort every paragraph's pPr children into the CT_PPr schema sequence.
    Run LAST, after every pass that touches paragraph properties: dedupe_ppr's
    merge and raw appends can leave children out of order, and an out-of-order
    property is silently dead in Word."""
    rank = {qn(f"w:{name}"): i for i, name in enumerate(_PPR_ORDER)}
    for ppr in doc.element.body.iter(qn("w:pPr")):
        children = list(ppr)
        if len(children) < 2:
            continue
        ordered = sorted(children, key=lambda c: rank.get(c.tag, len(rank)))
        if ordered != children:
            for c in children:
                ppr.remove(c)
            for c in ordered:
                ppr.append(c)


def apply_house_style(doc: Document, meta: dict | None = None) -> None:
    dedupe_ppr(doc)
    apply_page_geometry(doc)
    add_line_numbers(doc)
    apply_house_spacing_and_alignment(doc)
    apply_heading_look(doc)
    restyle_title(doc)
    if meta:
        rebuild_title_block(doc, meta)
        inject_keywords(doc, meta)
    restyle_abstract_heading(doc)
    bold_caption_leads(doc)
    fix_tables(doc)
    normalize_ppr_order(doc)


def _page_field() -> OxmlElement:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    return fld


def add_page_numbers(doc: Document, prefix: str = "") -> None:
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0]
        for run in list(para.runs):
            run._r.getparent().remove(run._r)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if prefix:
            para.add_run(prefix)
        para._p.append(_page_field())


def restart_page_numbering(doc: Document, start: int = 1) -> None:
    sect_pr = doc.sections[0]._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.insert_element_before(pg, *_SECT_PR_TAIL)
    pg.set(qn("w:start"), str(start))


def prepend_si_cover(doc: Document, cfg: dict, profile: dict, counts: dict) -> None:
    first = doc.paragraphs[0]
    lines = [
        "Supporting Information",
        f"Journal: {profile.get('journal', '')} — {cfg.get('ms_type', '')}",
        f"Contents: {counts.get('figures', 0)} figures, {counts.get('tables', 0)} tables",
        "Pages: (fill from final page count in Word — python-docx cannot paginate)",
    ]
    for text in lines:
        p = first.insert_paragraph_before(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brk = first.insert_paragraph_before("")
    run = brk.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


DEFAULT_TOC_ART_LABEL = "For Table of Contents Only"


def insert_toc_art(
    doc: Document,
    image_path: Path,
    width_mm: float | None = None,
    label: str = DEFAULT_TOC_ART_LABEL,
) -> None:
    """Append the TOC/abstract graphic at the END of the document.

    Per the official ACS "Guidelines for Table of Contents/Abstract
    Graphics" (dated 2024-02-28): the graphic is labeled "For Table of
    Contents Only" and placed on the LAST PAGE of the submitted manuscript
    (it may also be uploaded separately as "Graphics for manuscript").
    Produces, in order: a page break, a centered label paragraph, then a
    centered paragraph containing the image.
    """
    brk = doc.add_paragraph()
    run = brk.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)

    label_p = doc.add_paragraph(label)
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    kwargs = {"width": Mm(width_mm)} if width_mm else {}
    run.add_picture(str(image_path), **kwargs)


def quarto_render(project: Path, qmd: str, out_name: str, profile: dict) -> Path:
    cmd = ["quarto", "render", qmd, "--to", "docx", "--output", out_name]
    ref = profile.get("reference_doc")
    if ref:
        cmd += ["-M", f"reference-doc:{Path(profile['_dir']) / ref}"]
    csl = profile.get("csl")
    if csl:
        cmd += ["-M", f"csl:{Path(profile['_dir']) / csl}"]
    try:
        subprocess.run(cmd, cwd=project, check=True)
    except subprocess.CalledProcessError:
        raise SystemExit(f"quarto render failed for {qmd} (see quarto output above)")
    out_dir = project / "output"
    out_dir.mkdir(exist_ok=True)
    produced = project / out_name          # quarto writes next to the input
    target = out_dir / out_name
    if produced.exists() and produced != target:
        shutil.move(str(produced), target)
    if not target.exists():
        raise SystemExit(f"quarto render did not produce {target}")
    return target


def find_toc_art(project: Path) -> Path | None:
    for ext in ("png", "tiff", "tif", "jpg", "jpeg"):
        p = project / "figures" / f"toc-art.{ext}"
        if p.exists():
            return p
    return None


def postprocess_main(path: Path, profile: dict, cfg: dict, target: str, project: Path) -> None:
    doc = Document(str(path))
    if target == "submission":
        if profile.get("line_numbers"):
            add_line_numbers(doc)
        if profile.get("spacing") == "double":
            set_line_spacing(doc, 2.0)
        toc = (profile.get("toc_graphic") or {})
        if toc.get("required"):
            art = find_toc_art(project)
            if art is None:
                # Backstop only: main() gates this BEFORE quarto_render() runs
                # (see item A in quarto-docx-quirks.md / final-review-fixes),
                # so this branch should be unreachable in practice. Message
                # kept aligned with the pre-render gate's extension list.
                raise SystemExit(
                    "Profile requires TOC art but figures/toc-art.{png,tif,tiff,jpg,jpeg} is missing."
                )
            insert_toc_art(doc, art, toc.get("width_mm"), toc.get("label") or DEFAULT_TOC_ART_LABEL)
    set_fonts(doc, FONTS[target])
    apply_house_style(doc, read_front_matter(project, "index.qmd"))
    add_page_numbers(doc)
    doc.save(str(path))
    patch_theme_fonts(path, FONTS[target], track_changes=(target == "collab"))


def postprocess_si(path: Path, profile: dict, cfg: dict, target: str) -> None:
    doc = Document(str(path))
    si = profile.get("si") or {}
    set_fonts(doc, FONTS[target])
    apply_house_style(doc)  # no meta: SI keeps its plain author line, no keywords
    counts = {
        "figures": len(doc.inline_shapes),
        "tables": len(doc.tables),
    }
    if si.get("needs_cover_sheet"):
        prepend_si_cover(doc, cfg, profile, counts)
    add_page_numbers(doc, prefix=si.get("page_prefix", "S"))
    restart_page_numbering(doc, start=1)
    doc.save(str(path))
    patch_theme_fonts(path, FONTS[target], track_changes=(target == "collab"))


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--target", required=True, choices=("collab", "submission"))
    ap.add_argument("--project", default=".")
    args = ap.parse_args()
    project = Path(args.project).resolve()

    cfg = mslib.load_journal_config(project)
    profile = mslib.load_profile(cfg["journal"])
    mslib.manuscript_type(profile, cfg["ms_type"])  # fail fast on bad ms_type

    checks = validate.run_checks(project)
    validate.print_report(checks)
    hard = [c for c in checks if c.level == "HARD" and not c.ok]
    if hard and args.target == "submission":
        raise SystemExit("HARD checks failed — submission render refused (fix, or render --target collab).")

    # Gate BEFORE quarto ever runs: postprocess_main's TOC-art check used to
    # be the only enforcement, but it fires after quarto_render() has already
    # written output/main-submission.docx, leaving a half-processed file that
    # looks like a real deliverable when the run actually failed. Refusing
    # here means no docx is produced at all for this failure mode (backstop
    # check kept in postprocess_main in case this function is ever bypassed).
    if args.target == "submission" and (profile.get("toc_graphic") or {}).get("required"):
        if find_toc_art(project) is None:
            raise SystemExit(
                "Profile requires TOC art but figures/toc-art.{png,tif,tiff,jpg,jpeg} "
                "is missing — submission render refused."
            )

    main_docx = quarto_render(project, "index.qmd", f"main-{args.target}.docx", profile)
    postprocess_main(main_docx, profile, cfg, args.target, project)
    print(f"wrote {main_docx}")

    if (project / "si.qmd").exists():
        si_docx = quarto_render(project, "si.qmd", f"si-{args.target}.docx", profile)
        postprocess_si(si_docx, profile, cfg, args.target)
        print(f"wrote {si_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
