"""Regression tests for the OOXML pathology fixes.

These pin the behaviors documented in docs/docx-quirks.md — each test exists
because the corresponding bug shipped a wrong-looking manuscript at least once.
Scaffold phase: tests import the functions from legacy/render.py; when the
engine migrates into wongo.docxpatch, only the import below should change.

Run: uv run --with pytest --with python-docx --with pyyaml pytest
"""
from __future__ import annotations

import sys
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "legacy"))

from docx import Document  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

import wongo.docxpatch as render  # noqa: E402  (migrated engine module)
import wongo.styles as wstyles  # noqa: E402  (caption parsing moved here, step 2)


def _para_with_duplicate_ppr(doc):
    """Reproduce Quarto 1.10's invalid output: a paragraph with TWO pPr,
    the second carrying pStyle + jc=left (the one Word honors)."""
    p = doc.add_paragraph("Figure 1: A caption.")
    p.paragraph_format.line_spacing = 1.0  # creates the FIRST pPr
    first = p._p.find(qn("w:pPr"))
    extra = OxmlElement("w:pPr")
    st = OxmlElement("w:pStyle")
    st.set(qn("w:val"), "ImageCaption")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")
    extra.append(st)
    extra.append(jc)
    first.addnext(extra)  # the SECOND pPr, as Quarto 1.10 emits it
    return p


def test_dedupe_ppr_merges_duplicates_later_wins():
    doc = Document()
    p = _para_with_duplicate_ppr(doc)
    assert len([c for c in p._p if c.tag == qn("w:pPr")]) == 2
    render.dedupe_ppr(doc)
    pprs = [c for c in p._p if c.tag == qn("w:pPr")]
    assert len(pprs) == 1
    ppr = pprs[0]
    # later pPr's properties won; pStyle reordered to first child
    assert ppr[0].tag == qn("w:pStyle")
    jc = ppr.find(qn("w:jc"))
    assert jc is not None and jc.get(qn("w:val")) == "left"


def test_normalize_ppr_order_puts_spacing_before_jc():
    doc = Document()
    p = doc.add_paragraph("text")
    ppr = p._p.get_or_add_pPr()
    # deliberately wrong order: jc before spacing (Word silently drops jc)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "240")
    ppr.append(jc)
    ppr.append(spacing)
    render.normalize_ppr_order(doc)
    tags = [c.tag for c in p._p.find(qn("w:pPr"))]
    assert tags.index(qn("w:spacing")) < tags.index(qn("w:jc"))


def test_caption_lead_regex_handles_nbsp_and_si_prefix():
    m = wstyles._CAPTION_LEAD.match("Figure\xa01: A caption.")
    assert m and m.group(1).replace("\xa0", " ") == "Figure 1"
    m = wstyles._CAPTION_LEAD.match("Table S\xa02: An SI caption.")
    assert m is not None
    # prose like "Table 1 shows" (no delimiter) must NOT match the strict form
    assert wstyles._CAPTION_LEAD.match("Table 1 shows the result") is None


def test_style_profiles_load_and_kist_wcr_matches_legacy_constants():
    default = wstyles.load_style(None)
    assert default["_name"] == "default"
    kist = wstyles.load_style("kist-wcr")
    # the extracted profile must reproduce the HOUSE_* constants byte-for-byte
    # (page A4/25/25/30/25, double/single/justify lists, heading/title pt)
    assert kist["page"]["size_mm"] == [210, 297]
    assert kist["page"]["margins_mm"] == {"left": 25, "right": 25, "top": 30, "bottom": 25}
    assert kist["spacing"]["double"] == ["Normal", "Body Text", "First Paragraph", "Abstract"]
    assert kist["spacing"]["single"] == [
        "Compact", "Caption", "Image Caption", "Table Caption", "Footnote Text",
        "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title", "Author", "Date",
    ]
    assert kist["justify"] == [
        "Body Text", "First Paragraph", "Abstract", "Bibliography",
        "Caption", "Image Caption", "Table Caption",
        "Heading 1", "Heading 2", "Heading 3", "Heading 4",
    ]
    assert kist["headings_pt"] == {"Heading 1": 16, "Heading 2": 13, "Heading 3": 12, "Heading 4": 12}
    assert kist["title_pt"] == 14
    assert kist["title_block"]["style"] == "wr"
    assert kist["captions"]["lead"] == "bold" and kist["tables"]["rules"] == "booktabs"
    assert kist["keywords_line"] is True


def test_settings_compat_stamp(tmp_path):
    doc = Document()
    path = tmp_path / "t.docx"
    doc.save(str(path))
    render.patch_theme_fonts(path, "Cambria")
    import zipfile

    with zipfile.ZipFile(str(path)) as z:
        settings = z.read("word/settings.xml").decode()
        theme = z.read("word/theme/theme1.xml").decode()
    assert "compatibilityMode" in settings
    assert 'w:val="15"' in settings
    assert "Cambria" in theme


def test_track_changes_collab_only(tmp_path):
    for tracked in (True, False):
        doc = Document()
        path = tmp_path / f"t{tracked}.docx"
        doc.save(str(path))
        render.patch_theme_fonts(path, "Cambria", track_changes=tracked)
        import zipfile

        with zipfile.ZipFile(str(path)) as z:
            settings = z.read("word/settings.xml").decode()
        assert ("<w:trackChanges/>" in settings) is tracked
        if tracked and "<w:defaultTabStop" in settings:
            # schema order: trackChanges must precede defaultTabStop
            assert settings.index("<w:trackChanges/>") < settings.index("<w:defaultTabStop")


def test_patch_compat_mode_standalone(tmp_path):
    """The split-out compat stamp: inserts when absent, upgrades when < 15."""
    import re
    import zipfile

    for initial in (None, 'w:val="14"'):
        doc = Document()
        path = tmp_path / f"compat-{initial}.docx"
        doc.save(str(path))
        if initial is not None:
            with zipfile.ZipFile(str(path)) as z:
                items = {n: z.read(n) for n in z.namelist()}
            s = items["word/settings.xml"].decode()
            s = re.sub(
                r'<w:compatSetting[^>]*w:name="compatibilityMode"[^>]*/>',
                f'<w:compatSetting w:name="compatibilityMode" w:uri="x" {initial}/>',
                s,
            )
            items["word/settings.xml"] = s.encode()
            with zipfile.ZipFile(str(path), "w", zipfile.ZIP_DEFLATED) as z:
                for n, data in items.items():
                    z.writestr(n, data)
        render.patch_compat_mode(path)
        with zipfile.ZipFile(str(path)) as z:
            settings = z.read("word/settings.xml").decode()
        assert "compatibilityMode" in settings
        assert 'w:val="15"' in settings


def test_split_functions_match_combined_path(tmp_path):
    """patch_theme_fonts (combined) and manual font+compat application must
    produce identical settings.xml/theme bytes — the single-pass refactor
    must not change engine output."""
    import zipfile

    def read(path, name):
        with zipfile.ZipFile(str(path)) as z:
            return z.read(name)

    combined = tmp_path / "combined.docx"
    Document().save(str(combined))
    render.patch_theme_fonts(combined, "Cambria", track_changes=True)

    stepped = tmp_path / "stepped.docx"
    Document().save(str(stepped))
    render.patch_document_package(stepped, font="Cambria", track_changes=True)

    assert read(combined, "word/settings.xml") == read(stepped, "word/settings.xml")
    assert read(combined, "word/theme/theme1.xml") == read(stepped, "word/theme/theme1.xml")


def test_table_spacer_after_table_floats_only():
    """kist-wcr tables.spacer_after: empty single-spaced paragraph after TABLE
    floats, none after FIGURE floats (regression: NameError shipped because no
    test exercised fix_tables)."""
    from wongo.styles import fix_tables

    doc = Document()
    wrapper_t = doc.add_table(rows=1, cols=1)  # table float: nested data table
    inner = wrapper_t.rows[0].cells[0].add_table(rows=2, cols=2)
    assert inner is not None
    doc.add_paragraph("after table")
    style = {"tables": {"width": "full", "rules": "booktabs",
                        "bold_header_row": True, "bold_first_column": True,
                        "spacer_after": True}}
    fix_tables(doc, style)
    body = list(doc.element.body)
    ti = next(i for i, el in enumerate(body) if el.tag == qn("w:tbl"))
    nxt = body[ti + 1]
    assert nxt.tag == qn("w:p")
    assert not "".join(x.text or "" for x in nxt.iter(qn("w:t")))  # empty spacer
