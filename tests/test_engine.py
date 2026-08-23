"""Tests for the migrated render engine (HANDOFF step 4).

The SI cover sheet had TWO known gaps pinned here tests-first (see
HANDOFF-wongo-uplift.md): it printed journal+ms_type where the profile wants
AUTHORS (ES&T: cover sheet carries authors, title, page/figure/table counts),
and the page-count line shipped a literal placeholder string.
"""
from __future__ import annotations

import sys
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "legacy"))

from docx import Document

from wongo.engine import prepend_si_cover


def _meta():
    return {
        "title": "Electron Bifurcation in Wastewater",
        "author": [
            {"name": "Hoo Hugo Kim", "affiliations": [{"name": "KIST"}], "corresponding": True},
            {"name": "Jane Doe", "affiliations": [{"name": "KIST"}]},
        ],
    }


def test_cover_sheet_carries_authors_and_title_not_journal():
    doc = Document()
    doc.add_paragraph("SI body")
    prepend_si_cover(
        doc,
        profile={"journal": "Environmental Science & Technology", "si": {}},
        counts={"figures": 3, "tables": 2},
        meta=_meta(),
    )
    text = "\n".join(p.text for p in doc.paragraphs[:6])
    assert "Supporting Information" in text
    assert "Electron Bifurcation in Wastewater" in text
    assert "Hoo Hugo Kim" in text and "Jane Doe" in text
    # gap 1: journal name / ms_type must NOT appear anymore
    assert "Journal:" not in text
    assert "research-article" not in text


def test_cover_sheet_page_count_is_a_live_field_not_placeholder():
    doc = Document()
    doc.add_paragraph("SI body")
    prepend_si_cover(
        doc,
        profile={"si": {}},
        counts={"figures": 0, "tables": 0},
        meta=_meta(),
    )
    paras = doc.paragraphs[:6]
    texts = [p.text for p in paras]
    # gap 2: the literal placeholder must be gone
    assert not any("fill from final page count" in t for t in texts)
    # replaced by a NUMPAGES field Word resolves on open
    xml = paras[-1]._p.xml if paras else ""
    xml_all = "\n".join(p._p.xml for p in paras)
    assert "NUMPAGES" in xml_all
    assert any(t.startswith("Pages:") for t in texts)


def test_cover_sheet_ends_with_page_break_before_body():
    doc = Document()
    doc.add_paragraph("SI body")
    prepend_si_cover(doc, profile={"si": {}}, counts={"figures": 0, "tables": 0}, meta=_meta())
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    brs = list(doc._element.body.iter(f"{W}br"))
    assert any(br.get(f"{W}type") == "page" for br in brs)
