"""Tests for wongo.engine.diff — tracked-changes stamping for revisions.

The S6 promise: given the originally submitted DOCX and a revised render,
produce a copy of the revised document whose differences from the original
appear as real Word track changes (w:ins / w:del), so journals that demand
a marked-up revision get one without a manual Word Compare session.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from wongo.engine.diff import diff_documents


def _make_doc(tmp_path: Path, name: str, paragraphs: list[str]) -> Path:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    path = tmp_path / name
    doc.save(str(path))
    return path


def _body_paras(doc_path: Path) -> list:
    return Document(str(doc_path)).paragraphs


def _ins_texts(p) -> list[str]:
    out = []
    for ins in p._p.findall(qn("w:ins")):
        for r in ins.findall(qn("w:r")):
            for t in r.findall(qn("w:t")):
                out.append(t.text or "")
    return out


def _del_texts(p) -> list[str]:
    out = []
    for dele in p._p.findall(qn("w:del")):
        for r in dele.findall(qn("w:r")):
            for dt in r.findall(qn("w:delText")):
                out.append(dt.text or "")
    return out


def test_identical_documents_produce_no_changes(tmp_path):
    orig = _make_doc(tmp_path, "a.docx", ["Same one.", "Same two."])
    rev = _make_doc(tmp_path, "b.docx", ["Same one.", "Same two."])
    report = diff_documents(orig, rev, tmp_path / "out.docx")
    assert report["inserted_words"] == 0
    assert report["deleted_words"] == 0
    for p in _body_paras(tmp_path / "out.docx"):
        assert _ins_texts(p) == []
        assert _del_texts(p) == []


def test_word_replacement_emits_del_then_ins_in_same_paragraph(tmp_path):
    orig = _make_doc(tmp_path, "a.docx", ["The catalyst was copper."])
    rev = _make_doc(tmp_path, "b.docx", ["The catalyst was nickel."])
    diff_documents(orig, rev, tmp_path / "out.docx")
    p = _body_paras(tmp_path / "out.docx")[0]
    dels = "".join(_del_texts(p))
    inss = "".join(_ins_texts(p))
    assert "copper" in dels
    assert "nickel" in inss
    # unchanged words stay as plain runs outside any tracking element
    plain = "".join(
        t.text or ""
        for r in p._p.findall(qn("w:r"))
        for t in r.findall(qn("w:t"))
    )
    assert "catalyst" in plain and "copper" not in plain and "nickel" not in plain


def test_inserted_paragraph_is_fully_tracked(tmp_path):
    orig = _make_doc(tmp_path, "a.docx", ["First paragraph."])
    rev = _make_doc(tmp_path, "b.docx", ["First paragraph.", "A brand new paragraph."])
    diff_documents(orig, rev, tmp_path / "out.docx")
    paras = _body_paras(tmp_path / "out.docx")
    assert len(paras) == 2
    # python-docx's .text does not see inside w:ins; extract via the tracker
    assert "".join(_ins_texts(paras[1])) == "A brand new paragraph."
    assert paras[1].text == ""  # nothing outside the tracking wrapper


def test_deleted_paragraph_retained_as_full_deletion(tmp_path):
    orig = _make_doc(tmp_path, "a.docx", ["Keep me.", "Delete me entirely.", "Keep me too."])
    rev = _make_doc(tmp_path, "b.docx", ["Keep me.", "Keep me too."])
    diff_documents(orig, rev, tmp_path / "out.docx")
    paras = _body_paras(tmp_path / "out.docx")
    assert len(paras) == 3  # deletion retained in place
    deleted = ["".join(_del_texts(p)) for p in paras]
    assert "Delete me entirely." in deleted
    # the retained deletion paragraph carries no visible w:t text
    assert paras[1].text == ""


def test_deletions_use_del_text_element(tmp_path):
    orig = _make_doc(tmp_path, "a.docx", ["Original wording here."])
    rev = _make_doc(tmp_path, "b.docx", ["Revised wording here."])
    diff_documents(orig, rev, tmp_path / "out.docx")
    p = _body_paras(tmp_path / "out.docx")[0]
    for dele in p._p.findall(qn("w:del")):
        assert dele.findall(qn("w:r"))
        for r in dele.findall(qn("w:r")):
            assert r.findall(qn("w:delText"))
            assert not r.findall(qn("w:t"))  # deletions never use w:t


def test_author_and_date_stamped_on_tracking_elements(tmp_path):
    orig = _make_doc(tmp_path, "a.docx", ["Old line."])
    rev = _make_doc(tmp_path, "b.docx", ["New line."])
    diff_documents(orig, rev, tmp_path / "out.docx",
                   author="Coauthor Round", date="2026-08-25T00:00:00Z")
    p = _body_paras(tmp_path / "out.docx")[0]
    for tag in ("w:ins", "w:del"):
        el = p._p.find(qn(tag))
        assert el is not None
        assert el.get(qn("w:author")) == "Coauthor Round"
        assert el.get(qn("w:date")) == "2026-08-25T00:00:00Z"


def test_tables_are_left_untouched_but_reported(tmp_path):
    orig = Document(); orig.add_paragraph("Intro."); orig.add_table(rows=1, cols=2)
    orig.tables[0].cell(0, 0).text = "old"; orig.save(str(tmp_path / "a.docx"))
    rev = Document(); rev.add_paragraph("Intro."); rev.add_table(rows=1, cols=2)
    rev.tables[0].cell(0, 0).text = "new"; rev.save(str(tmp_path / "b.docx"))
    report = diff_documents(tmp_path / "a.docx", tmp_path / "b.docx",
                            tmp_path / "out.docx")
    assert report["tables_differ"] is True
    out = Document(str(tmp_path / "out.docx"))
    assert out.tables[0].cell(0, 0).text == "new"  # untouched, no crash
